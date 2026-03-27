"""
text_ingestor.py — Парсинг текстовых файлов в PerceptEvent.

Поддерживаемые форматы:
  .txt, .md   — встроенный Python (open)
  .pdf        — pymupdf (fitz), graceful fallback → plain text
  .docx       — python-docx, graceful fallback → plain text
  .json       — json stdlib (рекурсивное извлечение строк)
  .csv        — csv stdlib (строки → текст)

Стратегия чанкинга (paragraph-aware fixed chunking):
  1. Разбить документ на блоки по абзацам / пустым строкам / заголовкам
  2. Если блок > CHUNK_MAX_CHARS — разрезать по жёсткой длине с overlap
  3. Размер чанка: 1000–1500 символов, overlap: 100–150 символов

Каждый чанк → PerceptEvent(modality="text", source=..., content=..., ...)
"""

from __future__ import annotations

import csv
import json
import logging
import re
from pathlib import Path
from typing import Any, List, Optional

from brain.core.events import EventFactory, PerceptEvent
from brain.perception.metadata_extractor import MetadataExtractor
from brain.perception.validators import check_file_size, validate_file_path

_logger = logging.getLogger(__name__)

# ─── Опциональные зависимости ────────────────────────────────────────────────

try:
    import fitz  # pymupdf
    _FITZ_AVAILABLE = True
except ImportError:
    _FITZ_AVAILABLE = False
    _logger.debug("pymupdf (fitz) не установлен — PDF будет читаться как plain text")

try:
    import docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    _logger.debug("python-docx не установлен — DOCX будет читаться как plain text")

# ─── Константы чанкинга ──────────────────────────────────────────────────────

CHUNK_MIN_CHARS = 1000   # минимальный размер чанка
CHUNK_MAX_CHARS = 1500   # максимальный размер чанка (жёсткий лимит)
CHUNK_OVERLAP   = 120    # перекрытие между чанками

# Поддерживаемые расширения
TEXT_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".json", ".csv"}

# Паттерн заголовков Markdown
_MD_HEADER_RE = re.compile(r'^#{1,6}\s+', re.MULTILINE)


# ─── TextIngestor ─────────────────────────────────────────────────────────────

class TextIngestor:
    """
    Читает текстовые файлы и преобразует их в список PerceptEvent.

    Параметры:
        chunk_min:   минимальный размер чанка в символах (default: 1000)
        chunk_max:   максимальный размер чанка в символах (default: 1500)
        overlap:     перекрытие между чанками в символах (default: 120)
        extractor:   экземпляр MetadataExtractor (создаётся автоматически)

    Использование:
        ingestor = TextIngestor()
        events = ingestor.ingest("docs/нейрон.pdf")
        # → List[PerceptEvent]

        events = ingestor.ingest_text("Нейрон — это клетка...", source="user_input")
        # → List[PerceptEvent]
    """

    def __init__(
        self,
        chunk_min: int = CHUNK_MIN_CHARS,
        chunk_max: int = CHUNK_MAX_CHARS,
        overlap: int = CHUNK_OVERLAP,
        extractor: Optional[MetadataExtractor] = None,
    ):
        self._chunk_min = chunk_min
        self._chunk_max = chunk_max
        self._overlap = overlap
        self._extractor = extractor or MetadataExtractor()

    # ─── Публичный API ───────────────────────────────────────────────────────

    def ingest(
        self,
        file_path: str,
        session_id: str = "",
        trace_id: str = "",
    ) -> List[PerceptEvent]:
        """
        Прочитать файл и вернуть список PerceptEvent (по одному на чанк).

        Args:
            file_path:  путь к файлу
            session_id: ID сессии (прокидывается в события)
            trace_id:   ID трассировки

        Returns:
            List[PerceptEvent] — пустой список при ошибке или hard reject

        Raises:
            Не бросает исключений — все ошибки логируются.
        """
        path = Path(file_path)
        ext = path.suffix.lower()

        # B.2: Валидация пути (path traversal, null bytes, system dirs)
        safe, reason = validate_file_path(file_path)
        if not safe:
            _logger.warning("TextIngestor: файл отклонён валидацией — %s: %s", reason, file_path)
            return []

        if not path.exists():
            _logger.error("TextIngestor: файл не найден: %s", file_path)
            return []

        # B.2: Проверка размера файла
        size_ok, size_mb = check_file_size(file_path)
        if not size_ok:
            _logger.warning(
                "TextIngestor: файл слишком большой (%.1f MB): %s",
                size_mb, file_path,
            )
            return []

        if ext not in TEXT_EXTENSIONS:
            _logger.warning(
                "TextIngestor: неподдерживаемый формат '%s' для файла %s",
                ext, file_path,
            )
            return []

        _logger.info("TextIngestor: обработка файла %s (ext=%s)", file_path, ext)

        try:
            if ext == ".pdf":
                pages = self._read_pdf(file_path)
            elif ext == ".docx":
                pages = self._read_docx(file_path)
            elif ext == ".json":
                pages = self._read_json(file_path)
            elif ext == ".csv":
                pages = self._read_csv(file_path)
            else:  # .txt, .md
                pages = self._read_plain(file_path)
        except Exception as e:
            _logger.error("TextIngestor: ошибка чтения %s: %s", file_path, e)
            return []

        if not pages:
            _logger.warning("TextIngestor: файл пуст или текст не извлёкся: %s", file_path)
            return []

        events: List[PerceptEvent] = []
        chunk_counter = 0

        for page_num, page_text in pages:
            if not page_text or not page_text.strip():
                continue

            chunks = self._chunk_text(page_text)

            for chunk in chunks:
                reject, reason = MetadataExtractor.should_reject(chunk)
                if reject:
                    _logger.debug(
                        "TextIngestor: чанк отклонён (hard reject): %s — %s",
                        file_path, reason,
                    )
                    continue

                source_ref = f"{file_path}#p{page_num}" if page_num > 0 else file_path
                meta = self._extractor.extract(
                    text=chunk,
                    source=source_ref,
                    file_path=file_path,
                    page=page_num if page_num > 0 else None,
                    chunk_id=chunk_counter,
                )

                event = EventFactory.percept(
                    source=source_ref,
                    content=chunk,
                    modality="text",
                    quality=meta["quality"],
                    language=meta["language"],
                    session_id=session_id,
                    page=meta.get("page"),
                    chunk_id=chunk_counter,
                    file_size_kb=meta.get("file_size_kb", 0.0),
                    encoding=meta.get("encoding", "utf-8"),
                    quality_label=MetadataExtractor.quality_label(meta["quality"]),
                    warnings=meta.get("warnings", []),
                )
                if trace_id:
                    event.trace_id = trace_id

                events.append(event)
                chunk_counter += 1

        _logger.info(
            "TextIngestor: %s → %d чанков, %d событий",
            file_path, chunk_counter, len(events),
        )
        return events

    def ingest_text(
        self,
        text: str,
        source: str = "user_input",
        session_id: str = "",
        trace_id: str = "",
    ) -> List[PerceptEvent]:
        """
        Принять текст напрямую (не из файла) и вернуть список PerceptEvent.

        Args:
            text:       входной текст
            source:     идентификатор источника (default: "user_input")
            session_id: ID сессии
            trace_id:   ID трассировки

        Returns:
            List[PerceptEvent]
        """
        reject, reason = MetadataExtractor.should_reject(text)
        if reject:
            _logger.warning("TextIngestor.ingest_text: hard reject — %s", reason)
            return []

        chunks = self._chunk_text(text)
        events: List[PerceptEvent] = []

        for chunk_id, chunk in enumerate(chunks):
            reject, reason = MetadataExtractor.should_reject(chunk)
            if reject:
                continue

            meta = self._extractor.extract(
                text=chunk,
                source=source,
                chunk_id=chunk_id,
            )

            event = EventFactory.percept(
                source=source,
                content=chunk,
                modality="text",
                quality=meta["quality"],
                language=meta["language"],
                session_id=session_id,
                chunk_id=chunk_id,
                quality_label=MetadataExtractor.quality_label(meta["quality"]),
                warnings=meta.get("warnings", []),
            )
            if trace_id:
                event.trace_id = trace_id

            events.append(event)

        _logger.info(
            "TextIngestor.ingest_text: source='%s' → %d событий",
            source, len(events),
        )
        return events

    # ─── Чтение форматов ─────────────────────────────────────────────────────

    def _read_plain(self, file_path: str) -> List[tuple[int, str]]:
        """Читать .txt / .md файл."""
        encodings = ["utf-8", "utf-8-sig", "cp1251", "latin-1"]
        for enc in encodings:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    text = f.read()
                return [(0, text)]
            except (UnicodeDecodeError, LookupError):
                continue
        _logger.warning("TextIngestor: не удалось определить кодировку: %s", file_path)
        return []

    def _read_pdf(self, file_path: str) -> List[tuple[int, str]]:
        """Читать .pdf файл постранично."""
        if not _FITZ_AVAILABLE:
            _logger.warning(
                "TextIngestor: pymupdf не установлен, читаю PDF как plain text: %s",
                file_path,
            )
            return self._read_plain(file_path)

        pages: List[tuple[int, str]] = []
        try:
            doc = fitz.open(file_path)
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text("text")
                if text and text.strip():
                    pages.append((page_num, text))
            doc.close()
        except Exception as e:
            _logger.error("TextIngestor: ошибка чтения PDF %s: %s", file_path, e)
        return pages

    def _read_docx(self, file_path: str) -> List[tuple[int, str]]:
        """Читать .docx файл."""
        if not _DOCX_AVAILABLE:
            _logger.warning(
                "TextIngestor: python-docx не установлен, читаю DOCX как plain text: %s",
                file_path,
            )
            return self._read_plain(file_path)

        try:
            document = docx.Document(file_path)
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            # Таблицы
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)
            full_text = "\n\n".join(paragraphs)
            return [(0, full_text)] if full_text.strip() else []
        except Exception as e:
            _logger.error("TextIngestor: ошибка чтения DOCX %s: %s", file_path, e)
            return []

    def _read_json(self, file_path: str) -> List[tuple[int, str]]:
        """Читать .json файл — рекурсивно извлечь все строковые значения."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            strings = _extract_strings_from_json(data)
            text = "\n".join(strings)
            return [(0, text)] if text.strip() else []
        except Exception as e:
            _logger.error("TextIngestor: ошибка чтения JSON %s: %s", file_path, e)
            return []

    def _read_csv(self, file_path: str) -> List[tuple[int, str]]:
        """Читать .csv файл — строки → текст."""
        try:
            lines: List[str] = []
            with open(file_path, "r", encoding="utf-8", newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                    if row_text:
                        lines.append(row_text)
            text = "\n".join(lines)
            return [(0, text)] if text.strip() else []
        except Exception as e:
            _logger.error("TextIngestor: ошибка чтения CSV %s: %s", file_path, e)
            return []

    # ─── Чанкинг ─────────────────────────────────────────────────────────────

    def _chunk_text(self, text: str) -> List[str]:
        """
        Paragraph-aware fixed chunking.

        Алгоритм:
          1. Разбить текст на блоки по абзацам (пустые строки) и заголовкам
          2. Накапливать блоки в буфер до CHUNK_MAX_CHARS
          3. Если блок сам по себе > CHUNK_MAX_CHARS — разрезать жёстко с overlap
          4. Добавить overlap из конца предыдущего чанка в начало следующего

        Returns:
            List[str] — список чанков
        """
        if not text or not text.strip():
            return []

        # Шаг 1: разбить на параграфы
        paragraphs = _split_into_paragraphs(text)

        chunks: List[str] = []
        buffer = ""
        prev_tail = ""  # хвост предыдущего чанка для overlap

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Если параграф сам по себе слишком длинный — режем жёстко
            if len(para) > self._chunk_max:
                # Сначала сбрасываем буфер
                if buffer.strip():
                    chunk = (prev_tail + " " + buffer).strip() if prev_tail else buffer.strip()
                    chunks.append(chunk)
                    prev_tail = buffer[-self._overlap:] if len(buffer) > self._overlap else buffer
                    buffer = ""

                # Режем длинный параграф
                hard_chunks = _hard_split(para, self._chunk_max, self._overlap)
                for i, hc in enumerate(hard_chunks):
                    if i == 0 and prev_tail:
                        hc = (prev_tail + " " + hc).strip()
                    chunks.append(hc)
                prev_tail = hard_chunks[-1][-self._overlap:] if hard_chunks else ""
                continue

            # Если добавление параграфа превысит лимит — сбросить буфер
            candidate = (buffer + "\n\n" + para).strip() if buffer else para
            if len(candidate) > self._chunk_max and buffer.strip():
                chunk = (prev_tail + " " + buffer).strip() if prev_tail else buffer.strip()
                chunks.append(chunk)
                prev_tail = buffer[-self._overlap:] if len(buffer) > self._overlap else buffer
                buffer = para
            else:
                buffer = candidate

        # Остаток буфера
        if buffer.strip():
            chunk = (prev_tail + " " + buffer).strip() if prev_tail else buffer.strip()
            chunks.append(chunk)

        # Фильтруем слишком короткие чанки (< 10 символов)
        chunks = [c for c in chunks if len(c.strip()) >= 10]

        return chunks

    # ─── Статистика ──────────────────────────────────────────────────────────

    def supported_extensions(self) -> List[str]:
        """Список поддерживаемых расширений."""
        return sorted(TEXT_EXTENSIONS)

    def __repr__(self) -> str:
        return (
            f"TextIngestor(chunk={self._chunk_min}-{self._chunk_max}, "
            f"overlap={self._overlap}, "
            f"pdf={'yes' if _FITZ_AVAILABLE else 'fallback'}, "
            f"docx={'yes' if _DOCX_AVAILABLE else 'fallback'})"
        )


# ─── Вспомогательные функции ─────────────────────────────────────────────────

def _split_into_paragraphs(text: str) -> List[str]:
    """
    Разбить текст на параграфы по:
      - двойным переносам строк (пустые строки)
      - заголовкам Markdown (# ## ###)
      - горизонтальным разделителям (---, ===, ***)
    """
    # Нормализуем переносы строк
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Разбиваем по пустым строкам (2+ переноса)
    raw_blocks = re.split(r'\n{2,}', text)

    paragraphs: List[str] = []
    for block in raw_blocks:
        block = block.strip()
        if not block:
            continue

        # Если внутри блока есть заголовки Markdown — разбиваем по ним
        if _MD_HEADER_RE.search(block):
            sub_blocks = re.split(r'(?=^#{1,6}\s)', block, flags=re.MULTILINE)
            paragraphs.extend(sb.strip() for sb in sub_blocks if sb.strip())
        # Горизонтальные разделители
        elif re.match(r'^[-=*]{3,}\s*$', block):
            continue  # пропускаем разделители
        else:
            paragraphs.append(block)

    return paragraphs


def _hard_split(text: str, max_chars: int, overlap: int) -> List[str]:
    """
    Жёсткое разбиение длинного текста на чанки фиксированного размера с overlap.
    Пытается резать по границам предложений (. ! ?), иначе — по символам.
    """
    chunks: List[str] = []
    start = 0

    while start < len(text):
        end = min(start + max_chars, len(text))

        if end < len(text):
            # Ищем ближайшую границу предложения в последних 200 символах
            search_start = max(start, end - 200)
            segment = text[search_start:end]
            # Ищем последнюю точку/восклицание/вопрос
            match = None
            for m in re.finditer(r'[.!?]\s', segment):
                match = m
            if match:
                end = search_start + match.end()

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # Следующий старт с учётом overlap
        start = end - overlap if end - overlap > start else end

    return chunks


def _extract_strings_from_json(obj: Any, max_depth: int = 10) -> List[str]:
    """Рекурсивно извлечь все строковые значения из JSON-объекта."""
    if max_depth <= 0:
        return []
    if isinstance(obj, str):
        return [obj] if obj.strip() else []
    if isinstance(obj, dict):
        result: List[str] = []
        for k, v in obj.items():
            # Добавляем ключ как контекст если он информативный
            if isinstance(k, str) and len(k) > 2:
                result.append(f"{k}:")
            result.extend(_extract_strings_from_json(v, max_depth - 1))
        return result
    if isinstance(obj, (list, tuple)):
        result = []
        for item in obj:
            result.extend(_extract_strings_from_json(item, max_depth - 1))
        return result
    # Числа, bool, None — конвертируем в строку
    if obj is not None:
        return [str(obj)]
    return []
